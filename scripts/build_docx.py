"""
Convert medium_article.md to medium_article.docx with images and formatting.
Run from the docs/ folder: python build_docx.py
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCS_DIR = Path(__file__).parent
MD_FILE  = DOCS_DIR / "medium_article.md"
OUT_FILE = DOCS_DIR / "medium_article.docx"

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)


def add_image(path_str: str):
    img_path = DOCS_DIR / path_str.strip()
    if img_path.exists():
        doc.add_picture(str(img_path), width=Inches(5.5))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph(f"[Image not found: {path_str}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def apply_inline(run, text: str):
    """Apply bold/italic markdown to a run."""
    run.text = text


def add_paragraph_with_inline(para, line: str):
    """Parse bold (**text**) and italic (*text*) within a line."""
    pattern = re.compile(r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|(\[.+?\]\(.+?\)))')
    last = 0
    for m in pattern.finditer(line):
        if m.start() > last:
            para.add_run(line[last:m.start()])
        full = m.group(0)
        if full.startswith('***'):
            r = para.add_run(m.group(2))
            r.bold = r.italic = True
        elif full.startswith('**'):
            r = para.add_run(m.group(3))
            r.bold = True
        elif full.startswith('*'):
            r = para.add_run(m.group(4))
            r.italic = True
        elif full.startswith('`'):
            r = para.add_run(m.group(5))
            r.font.name = 'Courier New'
            r.font.size = Pt(10)
        elif full.startswith('['):
            # Markdown link — show display text, underlined
            link_text = re.match(r'\[(.+?)\]', full).group(1)
            r = para.add_run(link_text)
            r.underline = True
            r.font.color.rgb = RGBColor(0x1a, 0x0d, 0xab)
        last = m.end()
    if last < len(line):
        para.add_run(line[last:])


lines = MD_FILE.read_text(encoding="utf-8").splitlines()
i = 0

while i < len(lines):
    line = lines[i]

    # Blank line
    if not line.strip():
        i += 1
        continue

    # Image
    if line.strip().startswith("!["):
        m = re.match(r'!\[.*?\]\((.+?)\)', line.strip())
        if m:
            add_image(m.group(1))
        i += 1
        continue

    # Horizontal rule
    if line.strip().startswith("---"):
        doc.add_paragraph("─" * 60)
        i += 1
        continue

    # H1
    if line.startswith("# "):
        p = doc.add_heading(line[2:].strip(), level=1)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1
        continue

    # H2
    if line.startswith("## "):
        doc.add_heading(line[3:].strip(), level=2)
        i += 1
        continue

    # H3
    if line.startswith("### "):
        doc.add_heading(line[4:].strip(), level=3)
        i += 1
        continue

    # Italic metadata line (e.g. *Free tool · Singapore*)
    if line.startswith("*") and line.endswith("*") and not line.startswith("**"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line.strip("*"))
        r.italic = True
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        i += 1
        continue

    # Blockquote
    if line.startswith("> "):
        p = doc.add_paragraph(style="Intense Quote")
        add_paragraph_with_inline(p, line[2:].strip())
        i += 1
        continue

    # Bullet list
    if line.startswith("- ") or line.startswith("* "):
        p = doc.add_paragraph(style="List Bullet")
        add_paragraph_with_inline(p, line[2:].strip())
        i += 1
        continue

    # Numbered list
    m = re.match(r'^(\d+)\.\s+(.*)', line)
    if m:
        p = doc.add_paragraph(style="List Number")
        add_paragraph_with_inline(p, m.group(2).strip())
        i += 1
        continue

    # Normal paragraph
    p = doc.add_paragraph()
    add_paragraph_with_inline(p, line.strip())
    i += 1

doc.save(str(OUT_FILE))
print(f"Saved: {OUT_FILE}")
